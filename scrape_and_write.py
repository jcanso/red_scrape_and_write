import praw
import os
import docx
import re

# change this to something that makes sense on your machine
base_dir = "/home/jcanso/nosleep/"

# if dir does not exist, create it
if not os.path.isdir(base_dir):
    os.mkdir(base_dir)

# open a reddit instance
reddit = praw.Reddit(client_id='PUT_YOURS_HERE', client_secret="PUT_YOURS_HERE", user_agent="PUT_YOURS_HERE", username="PUT_YOURS_HERE", password="PUT_YOURS_HERE")

# get the subreddit and grab the latest 20 posts
subreddit = reddit.subreddit('nosleep')
latest_posts = subreddit.new(limit=20)

for post in latest_posts:
    """
    print(post.title)
    print("==========")
    print(post.selftext)
    print("\r\n")
    """

    # remove some characters from the title
    mutilated_title = re.sub('[\.\'\"]', '', post.title)

    # replace spaces with underscores, convert to ascii and back to get rid of weird characters for filenames
    file_name = base_dir + mutilated_title.replace(" ", "_") + ".docx"
    file_name = file_name.encode('ascii', 'ignore')
    file_name = file_name.decode()
    print(file_name)

    # open a document
    file = docx.Document()

    # populate with title and post text
    file.add_heading(post.title, 0)
    post_body = file.add_paragraph(post.selftext)

    # save document
    file.save(file_name)

